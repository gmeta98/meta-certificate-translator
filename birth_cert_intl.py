import os, re
from io import BytesIO
import boto3, streamlit as st
import zipfile
from dotenv import load_dotenv
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx import Document
from io import BytesIO
from datetime import datetime
from docx.shared import Cm
from docx.shared import Mm
import tempfile, os

# ── AWS / ENV ───────────────────────────────────────────────────────────────
load_dotenv()
textract = boto3.client(
    "textract",
    aws_access_key_id     = os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key = os.getenv("AWS_SECRET_ACCESS_KEY"),
    region_name           = os.getenv("AWS_REGION") or "us-east-2"
)

# ── STREAMLIT UI ────────────────────────────────────────────────────────────
st.set_page_config(page_title="Lindje Internacionale", layout="centered")
st.title("Certifikata Lindje Internacionale\nShqip - Italisht")
st.markdown("Ngarko nje ose me shume certifikata lindje internacionale dhe shkarko versionin italisht DOCX.")

uploaded_files = st.file_uploader("Ngarko certifikata", type=["pdf", "jpg", "jpeg", "png"], accept_multiple_files=True)
download_format = st.selectbox("Output format", ["Word (.docx)", "PDF (.pdf)"])

import streamlit as st

st.title("Perkthimi i Certifikates")

# --- Simple password gate (one shared password) ---
password = st.text_input("Password", type="password")
if "APP_PASSWORD" not in st.secrets:
    st.stop()  # safety if not configured
if password != st.secrets["APP_PASSWORD"]:
    st.warning("Enter the password to continue")
    st.stop()
# ---------------------------------------------------


# ── HELPER: Stato Civile from vertical checkboxes ───────────────────────────
def get_stato_from_vertical_boxes(blocks, bmap, gender=""):
    """
    Detects civil-status for the Albanian birth-certificate template.

    • Finds four status lines by keyword fragments:
        beqar   → single
        martu   → married
        shkuror → divorced
        vedov   → widowed
    • Finds the single handwritten “x”.
    • Chooses the label whose vertical centre is closest to the “x”.
    • Returns the gender-specific Italian form
        (celibe / nubile, coniugato / coniugata, divorziato / divorziata, vedovo / vedova).

    `gender` should be the already-extracted "Maschile" or “Femminile”.
    """

    fragments = [
        ("beqar",   0),   # index 0
        ("martu",   1),   # index 1
        ("shkuror", 2),   # index 2
        ("vedov",   3)    # index 3  (matches vedovo / vedova / vedov…)
    ]

    male   = ["Celibe",  "Coniugato",  "Divorziato", "Vedovo"]
    female = ["Nubile",  "Coniugata",  "Divorziata", "Vedova"]

    # ── 1. locate first WORD for each fragment and store its Y-centre
    centres = {}          # index → y
    for w in blocks:
        if w["BlockType"] != "WORD":
            continue
        text_low = w["Text"].strip().lower()
        for frag, idx in fragments:
            if frag in text_low and idx not in centres:
                bb = w["Geometry"]["BoundingBox"]
                centres[idx] = bb["Top"] + bb["Height"]/2

    if len(centres) < 4:
        return "[X] Stato non riconosciuto"

    # sort indices by Y
    ordered = sorted(centres.items(), key=lambda p: p[1])   # [(idx, y), …]

    # ── 2. find the handwritten “x”
    x_block = next(
        (w for w in blocks
         if w["BlockType"] == "WORD"
            and w["Text"].strip().lower() in ("x", "x.", "x,")),
        None
    )
    if not x_block:
        return "[X] Stato non riconosciuto"

    bbx = x_block["Geometry"]["BoundingBox"]
    x_centre_y = bbx["Top"] + bbx["Height"]/2

    # ── 3. choose the label whose centre-Y is nearest to the X
    best_idx, _ = min(ordered, key=lambda p: abs(p[1] - x_centre_y))

    # ── 4. return gender-specific Italian form
    g = (gender or "").lower()
    if g.startswith("f"):        # femminile
        return female[best_idx]
    if g.startswith("m"):        # maschile
        return male[best_idx]
    # fallback mixed form
    return f"{male[best_idx]} / {female[best_idx]}"


# __ HELPER: Word to PDF
def docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """
    Converts an in-memory DOCX (bytes) to PDF (bytes) using docx2pdf first.
    Falls back to LibreOffice on Linux.
    """
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "temp.docx")
        pdf_path  = os.path.join(tmp, "temp.pdf")

        # write the DOCX we already have
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        try:
            # primary: docx2pdf (needs Word on Win/macOS)
            from docx2pdf import convert
            convert(docx_path, pdf_path)       # ← create PDF
        except Exception:
            # fallback: LibreOffice (works headless everywhere)
            import subprocess, shlex
            cmd = f'libreoffice --headless --convert-to pdf --outdir "{tmp}" "{docx_path}"'
            subprocess.run(shlex.split(cmd), check=True)

        # read PDF back into memory
        with open(pdf_path, "rb") as f:
            return f.read()


# ── HELPER: Drop rotated / round-seal-stamp watermark LINE blocks ──────────
def filter_watermark_lines(blocks):
    """Strip OCR'd watermark text that comes from the round seal stamps on
    Albanian government certificates (e.g. the perimeter text on the
    'Ministria e Punëve të Brendshme' stamp). Only LINE blocks are
    removed — WORD blocks stay intact so CELL→WORD lookups still work."""
    SEAL_STAMP_WORDS = {
        "TIRANE", "TIRANÉ", "TIRANA",
        "BRENDSHME", "MINISTRIA", "PUNËVE", "PUNEVE",
    }
    out = []
    for b in blocks:
        if b.get("BlockType") != "LINE":
            out.append(b); continue
        bb = b.get("Geometry", {}).get("BoundingBox", {}) or {}
        w, h = bb.get("Width", 0), bb.get("Height", 0)
        txt = (b.get("Text") or "").strip()
        # 1. multi-char text whose bbox is taller than wide = rotated
        if w > 0 and len(txt) > 1 and h / w > 1.2:
            continue
        # 2. known stamp-perimeter words in the bottom area — only the
        # ALL-CAPS variant, so legitimate mixed-case prose ("...të
        # Brendshme") on the same line stays.
        if (bb.get("Top", 0) > 0.80
                and txt.isupper()
                and txt in SEAL_STAMP_WORDS):
            continue
        out.append(b)
    return out


# ── HELPER: Get Seal Block (last 2 lines) ────────────────────────────────────
def extract_seal_footer(blocks):
    import re

    lines = [b["Text"] for b in blocks if b["BlockType"] == "LINE"]

    # locate the *second* “Vulosur elektronikisht …”
    matches = [i for i, t in enumerate(lines)
               if "vulosur elektronikisht" in t.lower()]
    if len(matches) < 2:
        return ""                       # nothing found → bail out

    start = matches[1]
    date_line = ""
    hash_lines = []

    for raw in lines[start:]:
        txt = raw.strip()
        if not txt:
            continue

        # ── grab the date line ──────────────────────────────────────────────
        if not date_line and re.search(r"\d{4}/\d{2}/\d{2}", txt):
            # remove any leading "Date", "Datë", "Date:", "Datë:" etc.
            txt = re.sub(r"^(Date|Datë)\s*:?\s*", "", txt, flags=re.I).strip()
            date_line = f"In data {txt}"          # ← Italian label
            continue

        # ── grab hash / id lines after the date (hex or digits, ≥6 chars) ──
        if date_line and re.fullmatch(r"[A-Fa-f0-9]{6,}", txt):
            hash_lines.append(txt)
            continue

        # stop once we encounter anything else after collecting the seal
        if date_line and hash_lines:
            break

    if not date_line:
        return ""

    # Pick the Italian header based on which authority issued the seal
    seal_source = lines[start].lower()
    if "ministri" in seal_source:
        header = ["Timbro elettronico del Ministero", "degli Affari Interni"]
    else:
        header = ["Timbro elettronico della Direzione", "Generale dello Stato Civile"]

    parts = header + [date_line] + hash_lines
    return "\n".join(parts)


# ── HELPER: Albanian→Italian place exonyms ──────────────────────────────────
EXONYM_RULES = [
    (r"\bTiran[ëe]\b",   "Tirana"),
    (r"\bVlor[ëe]\b",    "Valona"),
    (r"\bDurr[ëe]s\b",   "Durazzo"),
    (r"\bShkod[ëe]r\b",  "Scutari"),
]

def map_exonyms(text: str) -> str:
    if not text:
        return text
    out = text
    for pat, repl in EXONYM_RULES:
        out = re.sub(pat, repl, out, flags=re.IGNORECASE)
    return out


# ── TABLE-FIELD EXTRACTION ──────────────────────────────────────────────────
def extract_table_fields(blocks, bmap):
    tbl = next((b for b in blocks if b["BlockType"] == "TABLE"), None)
    if not tbl: return {}

    rows = {}
    for rel in tbl.get("Relationships", []):
        if rel["Type"] != "CHILD": continue
        for cid in rel["Ids"]:
            cell = bmap[cid]
            if cell["BlockType"] != "CELL": continue
            r,c = cell["RowIndex"], cell["ColumnIndex"]
            txt = " ".join(
                w["Text"] for cr in cell.get("Relationships", [])
                for wid in cr.get("Ids", []) for w in (bmap[wid],)
                if w["BlockType"] == "WORD"
            ).strip()
            rows.setdefault(r, {})[c] = txt

    # ---------- ***NEW: clean up row 9 / col 2 abbreviations *** ----------
    res_raw = rows.get(9, {}).get(2, "")
    if res_raw:                                # only touch this one cell
        res_clean = (
            re.sub(r"Nd.",  "Ed.",  res_raw)   # Nd.  → Ed.
            .replace("H.",     "Int.")            # H.   → Int.
            .replace("Ap.",    "App.")             # Ap.  → App.
            .replace("Njësia",      "Sezione")
            .replace("Administrative",      "Amministrativa")
            .replace("NJËSIA",      "Sezione")
            .replace("ADMINISTRATIVE",      "Amministrativa")
            .replace("NJESIA",      "Sezione")
            .replace("Njesia",      "Sezione")
        )
        rows[9][2] = res_clean

    sesso_raw = rows.get(10, {}).get(2, "").strip().upper()
    if   sesso_raw == "M": sesso_val = "Maschile"
    elif sesso_raw == "F": sesso_val = "Femminile"
    else:                  sesso_val = sesso_raw


    result = {
        "Nome":              rows.get(2,  {}).get(2,""),
        "Cognome":           rows.get(3,  {}).get(2,""),
        "Numero personale":  rows.get(4,  {}).get(2,""),
        "Nome del padre":    rows.get(5,  {}).get(2,""),
        "Nome della madre":  rows.get(6,  {}).get(2,""),
        "Data di nascita":   rows.get(7,  {}).get(2,""),
        "Luogo di nascita":  rows.get(8,  {}).get(2,""),
        "Residenza":         rows.get(9,  {}).get(2,""),
        "Sesso":             sesso_val,
        "Stato Civile":      get_stato_from_vertical_boxes(blocks, bmap, sesso_val),
        "Cittadinanza":      rows.get(12, {}).get(2,""),
        "Cognome prima del matrimonio": rows.get(13, {}).get(2,""),
        "Data del rilascio": rows.get(14, {}).get(2,""),
        "ElectronicSeal":    extract_seal_footer(blocks),
    }

    # Normalize place names
    for k in ("Luogo di nascita", "Residenza"):
        result[k] = map_exonyms(result.get(k, ""))

    citt = (result.get("Cittadinanza") or "").strip().upper()
    if citt in ("ALB", "ALBANIA", "SHQIPTARE", "SHQIPTAR"):
        result["Cittadinanza"] = "Albanese"

    return result

# ── HEADER (Comune / Sezione) ───────────────────────────────────────────────
def extract_comune_sezione(blocks):
    lines = [b["Text"] for b in blocks if b["BlockType"] == "LINE"]
    comune = sezione = ""
    for i,l in enumerate(lines):
        if "Bashkia" in l:
            m = re.search(r"Bashkia\s+([A-ZÇËA-Za-zë\-]+)", l)
            if m: comune = m.group(1).title()
        if "Njësia Administrative" in l or "Njesia Administrative" in l:
            suf = l.split("Administrative",1)[1].strip()
            if suf.lower() in ("nr.","nr"):
                suf = suf + " " + (lines[i+1] if i+1 < len(lines) else "")
            sezione = suf.title()
    # normalize Comune
    comune = map_exonyms(comune)
    sezione = map_exonyms(sezione)

    return comune, sezione


# ── DOCX TEMPLATE ───────────────────────────────────────────────────────────
def make_docx(data):
    doc = Document()
    today = datetime.today().strftime("%d.%m.%Y")
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(1)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    

    # === Set base style: Times New Roman, black, size 11 ===
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    def add_paragraph(text, size=11, align="left", bold=False, italic=False, underline=False, indent_cm=0):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.color.rgb = RGBColor(0, 0, 0)

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

        if indent_cm > 0:
            p.paragraph_format.left_indent = Cm(indent_cm)

        p.alignment = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }.get(align, WD_PARAGRAPH_ALIGNMENT.LEFT)

        return p
    
    # === Single details table: title + rows ===
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.autofit = False  # we’ll control widths

    # Page inner width ≈ 21cm - 2cm - 2cm = 17cm
    left_w, right_w = Cm(9), Cm(8)  # tweak as you like; sum <= ~17cm
    tbl.columns[0].width = left_w
    tbl.columns[1].width = right_w

    # Some Word versions ignore column widths for new rows unless we reapply:
    def _set_row_widths(row):
        row.cells[0].width = left_w
        row.cells[1].width = right_w
   
   # ── Header row (Flag + REPUBBLICA | Comune + Sezione) ──
    row = tbl.add_row()
    _set_row_widths(row)

    # left header cell
    cell1 = row.cells[0]
    p1 = cell1.paragraphs[0]
    
    # add an empty line first
    p1.add_run("\n")

    img_path = os.path.join(os.getcwd(), "al_flag.png")
    if os.path.exists(img_path):
        r = p1.add_run()
        r.add_picture(img_path, width=Cm(0.7))
    r = p1.add_run("\n\nREPUBBLICA D'ALBANIA\n")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # right header cell
    cell2 = row.cells[1]
    p2 = cell2.paragraphs[0]
    lines = []
    if data.get("Comune"):
        lines.append(f"\n\nUfficio di Stato Civile Comune di {data['Comune']}")
    if data.get("Sezione"):
        lines.append(f"Sezione Amministrativa {data['Sezione']}")
    r = p2.add_run("\n".join(lines))
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # ── Title row merged across both columns ──
    row = tbl.add_row()
    _set_row_widths(row)
    merged = row.cells[0].merge(row.cells[1])
    p = merged.paragraphs[0]
    run = p.add_run("\nCERTIFICATO DI NASCITA\n")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    fields = [
        ("Nome",data["Nome"]), ("Cognome",data["Cognome"]),
        ("Numero personale",data["Numero personale"]),
        ("Nome del padre",data["Nome del padre"]),
        ("Nome della madre",data["Nome della madre"]),
        ("Data di nascita",data["Data di nascita"]),
        ("Luogo di nascita",data["Luogo di nascita"]),
        ("Residenza",data["Residenza"]),
        ("Sesso",data["Sesso"]),
        ("Stato Civile",data["Stato Civile"]),
        ("Cittadinanza",data["Cittadinanza"]),
        ("Cognome prima del matrimonio",data["Cognome prima del matrimonio"]),
        ("Data del rilascio", data["Data del rilascio"]),
        # 👇 Last row as a special marker
        ("\nTimbrato elettronicamente dalla Direzione Generale dello Stato Civile\n", None),
    ]

    for k, v in fields:
        if v is None:
            # special: merged last row
            row = tbl.add_row()
            _set_row_widths(row)
            merged = row.cells[0].merge(row.cells[1])

            para = merged.paragraphs[0]
            run = para.add_run(k)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

            merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after = Pt(5)
            para.paragraph_format.line_spacing = 1
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            continue

        # normal two-cell row
        row = tbl.add_row()
        _set_row_widths(row)
        cells = row.cells

        # vertical center left label only (or both if you prefer)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # left cell (label)
        para_left = cells[0].paragraphs[0]
        run_left = para_left.add_run(k)
        run_left.font.name = 'Times New Roman'
        run_left.font.size = Pt(11)
        run_left.font.color.rgb = RGBColor(0, 0, 0)
        para_left.paragraph_format.space_before = Pt(5)
        para_left.paragraph_format.space_after = Pt(5)
        para_left.paragraph_format.line_spacing = 1
        para_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # right cell (value)
        if k == "Cognome prima del matrimonio":
            display_val = v.strip() if v and v.strip() else "-------"
        else:
            display_val = v or ""
        para_right = cells[1].paragraphs[0]
        run_right = para_right.add_run(display_val)
        run_right.font.name = 'Times New Roman'
        run_right.font.size = Pt(11)
        run_right.font.color.rgb = RGBColor(0, 0, 0)
        para_right.paragraph_format.space_before = Pt(5)
        para_right.paragraph_format.space_after = Pt(5)
        para_right.paragraph_format.line_spacing = 1
        para_right.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

      

    # Electronic Seal
    if data.get("ElectronicSeal"):
        doc.add_paragraph()
        add_paragraph(data["ElectronicSeal"], size=10)

    # Footer sections (size 10)
    add_paragraph(
        "\nNota: Questo documento è stato generato e timbrato \nda una procedura automatica da un "
        "sistema elettronico \n(Direzione Generale di Stato Civile)\n", italic=True,
        size=10
    )
    
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False  # Disable Word's auto-resizing
    table.style = 'Table Grid'

    table.columns[0].width = Cm(11)
    table.rows[0].cells[0].width = Cm(11)  # Redundant but safer for compatibility

    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]  
    run = p.add_run(
        "Io, Vjollca META, traduttrice ufficiale della lingua italiana certificata dal Ministero "
        "della Giustizia con il numero di certificato 412 datato 31.07.2024, dichiaro di aver tradotto "
        "il testo presentatomi dalla lingua albanese all'italiano con precisione e responsabilità legale.\n"
        f"In data {today}."
)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    add_paragraph(
        "\n\nTraduzione eseguita da:\nVjollca META",
        size=11,
        align="center",
        indent_cm=11
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── MAIN FLOW ───────────────────────────────────────────────────────────────
if uploaded_files and st.button("Perkthe"):
    if len(uploaded_files) == 1:
        # Single file logic
        uploaded_file = uploaded_files[0]
        with st.spinner("Translating the certificate..."):
            resp = textract.analyze_document(
                Document={'Bytes': uploaded_file.read()},
                FeatureTypes=["FORMS", "TABLES", "LAYOUT"]
            )
            blocks = filter_watermark_lines(resp["Blocks"])
            bmap = {b["Id"]: b for b in blocks}

            data = extract_table_fields(blocks, bmap)
            data["Comune"], data["Sezione"] = extract_comune_sezione(blocks)

        with st.expander("🔍 Extracted Fields"): st.json(data)

        doc_buf = make_docx(data)

        today_str = datetime.today().strftime("%d-%m-%Y")
        nome = data.get("Nome", "Nome").strip().replace(" ", "_")
        cognome = data.get("Cognome", "Cognome").strip().replace(" ", "_")
        fname = f"{nome}_{cognome}_Certificato_di_Nascita_{today_str}"

        if download_format.startswith("PDF"):
            st.download_button("📥 Download PDF", doc_buf.getvalue(),
                               file_name=f"{fname}.pdf", mime="application/pdf")
        else:
            st.download_button("📥 Download DOCX", doc_buf.getvalue(),
                               file_name=f"{fname}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    else:
        # Multiple files logic → ZIP
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zipf:
            for uploaded_file in uploaded_files:
                with st.spinner(f"Translating {uploaded_file.name}..."):
                    resp = textract.analyze_document(
                        Document={'Bytes': uploaded_file.read()},
                        FeatureTypes=["FORMS", "TABLES", "LAYOUT"]
                    )
                    blocks = filter_watermark_lines(resp["Blocks"])
                    bmap = {b["Id"]: b for b in blocks}

                    data = extract_table_fields(blocks, bmap)
                    data["Comune"], data["Sezione"] = extract_comune_sezione(blocks)

                    doc_buf = make_docx(data)

                    today_str = datetime.today().strftime("%d-%m-%Y")
                    nome = data.get("Nome", "Nome").strip().replace(" ", "_")
                    cognome = data.get("Cognome", "Cognome").strip().replace(" ", "_")
                    fname = f"{nome}_{cognome}_Certificato_di_Nascita_{today_str}"

                    ext = "pdf" if download_format.startswith("PDF") else "docx"
                    zipf.writestr(f"{fname}.{ext}", doc_buf.getvalue())

        zip_buffer.seek(0)
        st.download_button(
            label="📥 Download All Translations (ZIP)",
            data=zip_buffer,
            file_name=f"certificati_tradotti_{datetime.today().strftime('%Y-%m-%d')}.zip",
            mime="application/zip"
        )
